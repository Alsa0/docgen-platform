# Service de génération de documents DOCX professionnels
# Utilise python-docx pour créer des documents Word formatés

import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config.settings import COMPANY_NAME, OUTPUT_DIR

logger = logging.getLogger(__name__)


def _add_header(doc, title: str, config: dict):
    """Ajoute un en-tête professionnel au document."""
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.text = f"{COMPANY_NAME} — {title}"
    p.style.font.size = Pt(9)
    p.style.font.color.rgb = RGBColor(100, 100, 100)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_footer(doc, reference: str = ""):
    """Ajoute un pied de page avec numéro et référence."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = f"{COMPANY_NAME} | {reference} | Confidentiel"
    p.style.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title_page(doc, title: str, subtitle: str, config: dict):
    """Ajoute une page de garde."""
    for _ in range(4):
        doc.add_paragraph("")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(COMPANY_NAME)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.bold = True

    doc.add_paragraph("")
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = t2.add_run(title)
    run2.font.size = Pt(28)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0, 51, 102)

    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = t3.add_run(subtitle)
    run3.font.size = Pt(14)
    run3.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f"Client : {config.get('client_nom', 'N/A')}\n").font.size = Pt(12)
    info.add_run(f"Date : {datetime.now().strftime('%d/%m/%Y')}\n").font.size = Pt(12)
    if config.get("version"):
        info.add_run(f"Version : {config.get('version', '1.0')}").font.size = Pt(12)
    doc.add_page_break()


def _add_section_title(doc, number: int, title: str, level: int = 1):
    """Ajoute un titre de section numéroté."""
    heading = doc.add_heading(f"{number}. {title}", level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def _add_table(doc, headers: list, rows: list, total_row: list = None):
    """Crée un tableau formaté avec en-têtes et données."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-têtes
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Données
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val) if val is not None else ""
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    # Ligne total
    if total_row:
        row = table.add_row()
        for i, val in enumerate(total_row):
            row.cells[i].text = str(val) if val is not None else ""
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
    doc.add_paragraph("")
    return table


def generate_bom_docx(config: dict, bom_items: list) -> str:
    """Génère un document DOCX pour le Bill of Materials."""
    try:
        doc = Document()
        devise = config.get("devise", "MGA")
        _add_header(doc, "BILL OF MATERIALS", config)
        _add_footer(doc, config.get("reference_projet", "BOM"))

        # Titre
        t = doc.add_heading("BILL OF MATERIALS", level=0)
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Infos projet
        doc.add_paragraph(f"Client : {config.get('client_nom', 'N/A')}")
        doc.add_paragraph(f"Projet : {config.get('projet_titre', 'N/A')}")
        doc.add_paragraph(f"Date : {datetime.now().strftime('%d/%m/%Y')}")
        doc.add_paragraph(f"Référence : {config.get('reference_projet', 'N/A')}")
        doc.add_paragraph(f"Type : {config.get('type_projet', 'N/A')}")
        doc.add_paragraph("")

        # Tableau principal
        headers = ["N°", "Référence", "Désignation", "Marque/Modèle", "Qté",
                    f"Prix Unit. ({devise})", f"Prix Total ({devise})", "Fournisseur"]
        rows = []
        total = 0
        for i, item in enumerate(bom_items):
            qty = item.get("quantite", 1)
            prix_u = item.get("prix_unitaire", 0)
            prix_t = qty * prix_u
            total += prix_t
            rows.append([
                i + 1, item.get("reference", ""),
                item.get("designation", ""), f"{item.get('marque', '')} {item.get('modele', '')}",
                qty, f"{prix_u:,.0f}", f"{prix_t:,.0f}", item.get("fournisseur", "")
            ])
        total_row = ["", "", "", "", "", "TOTAL", f"{total:,.0f}", ""]
        _add_table(doc, headers, rows, total_row)

        # Sources et liens
        _add_section_title(doc, 1, "Sources et liens")
        for item in bom_items:
            url = item.get("url", "")
            if url:
                doc.add_paragraph(f"• {item.get('designation', 'N/A')} : {url}", style="List Bullet")

        # Validité
        doc.add_paragraph("")
        validite = config.get("validite_jours", 30)
        doc.add_paragraph(f"Cette offre est valide {validite} jours à compter de la date d'émission.")
        doc.add_paragraph("Les prix sont indicatifs et peuvent varier selon la disponibilité.")

        filename = f"BOM_{config.get('client_nom', 'client')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logger.info(f"BOM DOCX généré : {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Erreur génération BOM DOCX : {e}")
        raise


def generate_sow_docx(config: dict, sow_content: dict, bom_items: list = None) -> str:
    """Génère un document DOCX pour le Scope of Work."""
    try:
        doc = Document()
        devise = config.get("devise", "MGA")
        _add_header(doc, "SCOPE OF WORK", config)
        _add_footer(doc, config.get("reference_projet", "SOW"))
        _add_title_page(doc, "SCOPE OF WORK", config.get("projet_titre", ""), config)

        # 1. Overview
        _add_section_title(doc, 1, "Overview")
        doc.add_paragraph(sow_content.get("overview", ""))

        # 2. Objectifs
        _add_section_title(doc, 2, "Objectifs")
        for obj in sow_content.get("objectifs", []):
            doc.add_paragraph(obj, style="List Bullet")

        # 3. Périmètre inclus
        _add_section_title(doc, 3, "Périmètre inclus")
        for item in sow_content.get("perimetre_inclus", []):
            doc.add_paragraph(item, style="List Bullet")

        # 4. Périmètre exclus
        _add_section_title(doc, 4, "Périmètre exclus")
        for item in sow_content.get("perimetre_exclus", config.get("exclusions", [])):
            doc.add_paragraph(item, style="List Bullet")

        # 5. Tâches détaillées
        _add_section_title(doc, 5, "Tâches détaillées")
        taches = sow_content.get("taches", [])
        headers_t = ["N°", "Tâche", "Responsable", "Durée", "Livrables"]
        rows_t = []
        for t in taches:
            rows_t.append([
                t.get("numero", ""), t.get("titre", ""),
                t.get("responsable", ""), f"{t.get('duree_jours', '')} jours",
                ", ".join(t.get("livrables", []))
            ])
        _add_table(doc, headers_t, rows_t)
        for t in taches:
            doc.add_heading(f"Tâche {t.get('numero', '')}: {t.get('titre', '')}", level=3)
            doc.add_paragraph(t.get("description", ""))

        # 6. Équipements (BOM)
        if bom_items:
            _add_section_title(doc, 6, "Équipements")
            headers_b = ["N°", "Désignation", "Qté", f"Prix ({devise})"]
            rows_b = []
            total = 0
            for i, item in enumerate(bom_items):
                pt = item.get("quantite", 1) * item.get("prix_unitaire", 0)
                total += pt
                rows_b.append([i+1, item.get("designation", ""), item.get("quantite", 1), f"{pt:,.0f}"])
            _add_table(doc, headers_b, rows_b, ["", "TOTAL", "", f"{total:,.0f}"])

        # 7. Jalons
        section_num = 7 if bom_items else 6
        _add_section_title(doc, section_num, "Jalons et planning")
        jalons = sow_content.get("jalons", [])
        if jalons:
            h_j = ["N°", "Jalon", "Date prévue", "Livrables"]
            r_j = [[j.get("numero",""), j.get("titre",""), j.get("date_prevue",""),
                     ", ".join(j.get("livrables", []))] for j in jalons]
            _add_table(doc, h_j, r_j)

        # 8. Critères d'acceptation
        _add_section_title(doc, section_num+1, "Critères d'acceptation")
        for c in config.get("criteres_acceptation", []):
            doc.add_paragraph(c, style="List Bullet")

        # 9. Conditions de paiement
        _add_section_title(doc, section_num+2, "Conditions de paiement")
        doc.add_paragraph(sow_content.get("conditions_paiement", "À définir"))

        # 10. Hypothèses
        _add_section_title(doc, section_num+3, "Hypothèses et dépendances")
        for h in sow_content.get("hypotheses", []):
            doc.add_paragraph(h, style="List Bullet")

        filename = f"SOW_{config.get('client_nom', 'client')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logger.info(f"SOW DOCX généré : {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Erreur génération SOW DOCX : {e}")
        raise


def generate_ot_docx(config: dict, ot_content: dict, sow_summary: dict = None, bom_summary: list = None) -> str:
    """Génère un document DOCX pour l'Offre Technique."""
    try:
        doc = Document()
        devise = config.get("devise", "MGA")
        _add_header(doc, "OFFRE TECHNIQUE", config)
        _add_footer(doc, config.get("reference_offre", "OT"))
        _add_title_page(doc, "OFFRE TECHNIQUE", config.get("projet_titre", ""), config)

        # 1. Présentation entreprise
        _add_section_title(doc, 1, "Notre entreprise")
        doc.add_paragraph(ot_content.get("presentation_entreprise", ""))

        # 2. Compréhension du besoin
        _add_section_title(doc, 2, "Compréhension du besoin")
        doc.add_paragraph(ot_content.get("comprehension_besoin", ""))

        # 3. Solution proposée
        _add_section_title(doc, 3, "Solution proposée")
        doc.add_paragraph(ot_content.get("solution_proposee", ""))

        # 4. Méthodologie
        _add_section_title(doc, 4, "Méthodologie")
        doc.add_paragraph(ot_content.get("methodologie", ""))

        # 5. Planning
        _add_section_title(doc, 5, "Planning prévisionnel")
        planning = ot_content.get("planning", [])
        if planning:
            h = ["Phase", "Durée", "Description"]
            r = [[p.get("phase",""), p.get("duree",""), p.get("description","")] for p in planning]
            _add_table(doc, h, r)

        # 6. Équipe
        _add_section_title(doc, 6, "Notre équipe")
        equipe = ot_content.get("equipe_projet", [])
        if equipe:
            for membre in equipe:
                doc.add_paragraph(f"• {membre.get('nom_role', '')} : {membre.get('responsabilites', '')}")

        # 7. Budget
        _add_section_title(doc, 7, "Budget global")
        budget_detail = ot_content.get("budget_detail", [])
        if budget_detail:
            h_b = ["Poste", f"Montant ({devise})", "Description"]
            r_b = [[b.get("poste",""), f"{b.get('montant',0):,.0f}", b.get("description","")] for b in budget_detail]
            _add_table(doc, h_b, r_b)

        # 8. Conditions
        _add_section_title(doc, 8, "Conditions et validité")
        doc.add_paragraph(ot_content.get("conditions_validite", f"Offre valide {config.get('validite_jours', 30)} jours."))
        doc.add_paragraph(ot_content.get("conditions_paiement", ""))

        # Annexe SOW
        if sow_summary:
            doc.add_page_break()
            doc.add_heading("ANNEXE A — Résumé du Scope of Work", level=1)
            doc.add_paragraph(sow_summary.get("overview", ""))
            for t in sow_summary.get("taches", [])[:5]:
                doc.add_paragraph(f"• {t.get('titre', '')}: {t.get('description', '')[:100]}...", style="List Bullet")

        # Annexe BOM
        if bom_summary:
            doc.add_page_break()
            doc.add_heading("ANNEXE B — Résumé du Bill of Materials", level=1)
            h_bom = ["Désignation", "Qté", f"Prix ({devise})"]
            r_bom = [[b.get("designation",""), b.get("quantite",1),
                       f"{b.get('quantite',1)*b.get('prix_unitaire',0):,.0f}"] for b in bom_summary]
            total = sum(b.get("quantite",1)*b.get("prix_unitaire",0) for b in bom_summary)
            _add_table(doc, h_bom, r_bom, ["TOTAL", "", f"{total:,.0f}"])

        filename = f"OT_{config.get('client_nom', 'client')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logger.info(f"OT DOCX généré : {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Erreur génération OT DOCX : {e}")
        raise


def generate_ir_docx(config: dict, ir_content: dict) -> str:
    """Génère un document DOCX pour le Rapport d'Intervention."""
    try:
        doc = Document()
        _add_header(doc, "RAPPORT D'INTERVENTION", config)
        _add_footer(doc, config.get("reference_ticket", "IR"))

        t = doc.add_heading("RAPPORT D'INTERVENTION", level=0)
        for run in t.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)

        # Statut bandeau
        statut = config.get("statut", "terminé")
        statut_colors = {"terminé": RGBColor(0,128,0), "en cours": RGBColor(255,165,0), "action requise": RGBColor(200,0,0)}
        sp = doc.add_paragraph()
        sr = sp.add_run(f"  STATUT : {statut.upper()}  ")
        sr.bold = True
        sr.font.size = Pt(14)
        sr.font.color.rgb = statut_colors.get(statut, RGBColor(0,0,0))

        # 1. Informations
        _add_section_title(doc, 1, "Informations de l'intervention")
        info_data = [
            ("Client", config.get("client_nom", "")),
            ("Site", config.get("site_intervention", "")),
            ("Date", config.get("date_intervention", "")),
            ("Horaires", f"{config.get('heure_debut', '')} - {config.get('heure_fin', '')}"),
            ("Techniciens", ", ".join(config.get("techniciens", []))),
            ("Type", config.get("type_intervention", "")),
            ("Référence", config.get("reference_ticket", "")),
        ]
        table_info = doc.add_table(rows=len(info_data), cols=2)
        table_info.style = "Light Grid Accent 1"
        for i, (label, val) in enumerate(info_data):
            table_info.rows[i].cells[0].text = label
            table_info.rows[i].cells[1].text = str(val)
            for p in table_info.rows[i].cells[0].paragraphs:
                for run in p.runs:
                    run.bold = True
        doc.add_paragraph("")

        # 2. Mission
        _add_section_title(doc, 2, "Description de la mission")
        doc.add_paragraph(config.get("description_mission", ""))
        doc.add_paragraph(ir_content.get("resume_intervention", ""))

        # 3. Travaux réalisés
        _add_section_title(doc, 3, "Travaux réalisés")
        travaux = ir_content.get("travaux_details", [])
        if travaux:
            h_t = ["N°", "Tâche", "Statut", "Observations"]
            r_t = [[t.get("numero",""), t.get("tache",""), t.get("statut",""), t.get("observations","")] for t in travaux]
            _add_table(doc, h_t, r_t)

        # 4. Observations
        _add_section_title(doc, 4, "Observations techniques")
        doc.add_paragraph(config.get("observations", ir_content.get("analyse_problemes", "")))

        # 5. Problèmes et solutions
        _add_section_title(doc, 5, "Problèmes rencontrés et solutions")
        doc.add_paragraph(f"Problèmes : {config.get('problemes_rencontres', 'Aucun')}")
        doc.add_paragraph(f"Solutions : {ir_content.get('solutions_detaillees', config.get('solutions_appliquees', 'N/A'))}")

        # 6. Recommandations
        _add_section_title(doc, 6, "Recommandations")
        for rec in ir_content.get("recommandations", []):
            doc.add_paragraph(rec, style="List Bullet")

        # 7. Actions requises
        if statut == "action requise":
            _add_section_title(doc, 7, "Actions requises")
            for act in ir_content.get("actions_requises", []):
                doc.add_paragraph(
                    f"• [{act.get('priorite', 'moyenne')}] {act.get('action', '')} "
                    f"— {act.get('responsable', '')} — Échéance : {act.get('echeance', 'N/A')}"
                )

        # 8. Conclusion
        _add_section_title(doc, 8, "Conclusion")
        doc.add_paragraph(ir_content.get("conclusion", ""))

        # Pièces utilisées
        pieces = config.get("pieces_utilisees", [])
        if pieces:
            _add_section_title(doc, 9, "Pièces et équipements utilisés")
            for p in pieces:
                doc.add_paragraph(p, style="List Bullet")

        filename = f"IR_{config.get('client_nom', 'client')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logger.info(f"IR DOCX généré : {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Erreur génération IR DOCX : {e}")
        raise


def generate_lld_docx(config: dict, lld_content: dict, bom_items: list = None) -> str:
    """Génère un document DOCX pour le Low Level Design."""
    try:
        doc = Document()
        devise = config.get("devise", "MGA")
        _add_header(doc, "LOW LEVEL DESIGN", config)
        _add_footer(doc, config.get("version", "LLD"))
        _add_title_page(doc, "LOW LEVEL DESIGN", config.get("projet_titre", ""), config)

        # 1. Contexte
        _add_section_title(doc, 1, "Contexte et objectifs")
        doc.add_paragraph(lld_content.get("contexte_objectifs", ""))

        # 2. Architecture
        _add_section_title(doc, 2, "Architecture générale")
        doc.add_paragraph(lld_content.get("architecture_generale", ""))

        # 3. Spécifications équipements
        _add_section_title(doc, 3, "Spécifications des équipements")
        specs = lld_content.get("specifications_equipements", [])
        if specs:
            for spec in specs:
                doc.add_heading(f"{spec.get('nom', '')} — {spec.get('marque_modele', '')}", level=3)
                doc.add_paragraph(f"Rôle : {spec.get('role', '')}")
                doc.add_paragraph(f"Spécifications : {spec.get('specs_techniques', '')}")
                doc.add_paragraph(f"Configuration : {spec.get('configuration_principale', '')}")
                if spec.get("datasheet_url"):
                    doc.add_paragraph(f"Datasheet : {spec['datasheet_url']}")

        # 4. Plan d'adressage
        _add_section_title(doc, 4, "Plan d'adressage IP")
        plan = lld_content.get("plan_adressage", [])
        if plan:
            h_ip = ["VLAN", "Nom", "Réseau", "Passerelle", "DHCP", "Usage"]
            r_ip = [[p.get("vlan_id",""), p.get("nom",""), p.get("reseau",""),
                      p.get("passerelle",""), p.get("plage_dhcp",""), p.get("usage","")] for p in plan]
            _add_table(doc, h_ip, r_ip)

        # 5. Schéma d'interconnexion
        _add_section_title(doc, 5, "Schéma d'interconnexion")
        doc.add_paragraph(lld_content.get("schema_interconnexion", ""))

        # 6. Configurations
        _add_section_title(doc, 6, "Configurations détaillées")
        configs_detail = lld_content.get("configurations_detaillees", [])
        for cfg in configs_detail:
            doc.add_heading(cfg.get("equipement", ""), level=3)
            p = doc.add_paragraph()
            run = p.add_run(cfg.get("configuration", ""))
            run.font.name = "Consolas"
            run.font.size = Pt(9)

        # 7. Procédures
        _add_section_title(doc, 7, "Procédures d'installation et de test")
        procedures = lld_content.get("procedures_installation", [])
        if procedures:
            h_p = ["Étape", "Titre", "Description", "Validation"]
            r_p = [[p.get("etape",""), p.get("titre",""), p.get("description",""), p.get("validation","")] for p in procedures]
            _add_table(doc, h_p, r_p)

        # 8. Références
        _add_section_title(doc, 8, "Références et datasheets")
        refs = lld_content.get("references_datasheets", [])
        for ref in refs:
            doc.add_paragraph(f"• {ref.get('equipement', '')} : {ref.get('url', '')} ({ref.get('source', '')})")

        # BOM si inclus
        if bom_items:
            doc.add_page_break()
            doc.add_heading("ANNEXE — Bill of Materials", level=1)
            h_b = ["N°", "Désignation", "Marque/Modèle", "Qté", f"Prix ({devise})"]
            r_b = []
            total = 0
            for i, item in enumerate(bom_items):
                pt = item.get("quantite", 1) * item.get("prix_unitaire", 0)
                total += pt
                r_b.append([i+1, item.get("designation",""),
                            f"{item.get('marque','')} {item.get('modele','')}", item.get("quantite",1), f"{pt:,.0f}"])
            _add_table(doc, h_b, r_b, ["", "", "TOTAL", "", f"{total:,.0f}"])

        filename = f"LLD_{config.get('client_nom', 'client')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(OUTPUT_DIR, filename)
        doc.save(filepath)
        logger.info(f"LLD DOCX généré : {filepath}")
        return filepath
    except Exception as e:
        logger.error(f"Erreur génération LLD DOCX : {e}")
        raise
